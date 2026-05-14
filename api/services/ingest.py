"""
FinHouse — File Ingest Service
Parse documents → chunk → embed via BGE-M3 → upsert into Milvus.

v2 additions (gated behind settings.RAG_* flags, fully backward compatible):
- New collection `finhouse_chunks_v2` with HNSW dense index + sparse field
- Hybrid retrieval: dense + BGE-M3 sparse (lexical) fused via RRF
- HyDE: retrieve_context accepts multiple query strings (rewritten + hypos)
- Semantic chunking: embedding-based topic boundary detection (opt-in)

Rollback: set settings.RAG_COLLECTION="finhouse_chunks" — the legacy
IVF_FLAT dense-only path is preserved end-to-end and never deleted.
"""

import io
import os
import re
import hashlib
import logging
import threading
from datetime import datetime, timezone
from typing import Optional
from uuid import UUID

import httpx
from minio import Minio

from config import get_settings

settings = get_settings()
logger = logging.getLogger("finhouse.ingest")

# ── Supported formats ───────────────────────────────────────
SUPPORTED_EXTENSIONS = {"pdf", "md", "txt", "docx"}

# ── Limits ──────────────────────────────────────────────────
MAX_FILE_SIZE_MB = 200          # reject files larger than this
MAX_CHUNKS_PER_FILE = 3000      # safety cap

# ── Chunking config ─────────────────────────────────────────
CHUNK_CHARS = 1500              # ~512 tokens per chunk
CHUNK_OVERLAP = 200             # overlap in chars
MIN_CHUNK_LENGTH = 50           # discard chunks shorter than this


# ════════════════════════════════════════════════════════════
# MinIO helpers
# ════════════════════════════════════════════════════════════

_minio_client: Optional[Minio] = None


def get_minio_client() -> Minio:
    """Return a reusable MinIO client (connection pooled internally by urllib3)."""
    global _minio_client
    if _minio_client is None:
        _minio_client = Minio(
            f"{settings.MINIO_HOST}:{settings.MINIO_PORT}",
            access_key=settings.MINIO_ROOT_USER,
            secret_key=settings.MINIO_ROOT_PASSWORD,
            secure=False,
        )
    return _minio_client


def ensure_bucket(client: Minio):
    try:
        if not client.bucket_exists(settings.MINIO_BUCKET):
            client.make_bucket(settings.MINIO_BUCKET)
    except Exception as e:
        logger.warning(f"Bucket check/create failed (may already exist): {e}")


def upload_to_minio(
    content: bytes,
    object_name: str,
    content_type: str = "application/octet-stream",
):
    client = get_minio_client()
    ensure_bucket(client)
    client.put_object(
        settings.MINIO_BUCKET,
        object_name,
        io.BytesIO(content),
        len(content),
        content_type=content_type,
    )


def download_from_minio(object_name: str) -> bytes:
    """Download with proper resource cleanup even on read errors."""
    client = get_minio_client()
    resp = None
    try:
        resp = client.get_object(settings.MINIO_BUCKET, object_name)
        data = resp.read()
        return data
    finally:
        if resp is not None:
            resp.close()
            resp.release_conn()


def delete_file_object(bucket: str, object_name: str):
    """Remove an object from MinIO. Idempotent — missing objects don't raise."""
    client = get_minio_client()
    try:
        client.remove_object(bucket, object_name)
    except Exception as e:
        # minio-py raises S3Error on missing object; we tolerate that
        err_code = getattr(e, "code", None)
        if err_code in ("NoSuchKey", "NoSuchBucket"):
            return
        raise


# ════════════════════════════════════════════════════════════
# Document Parsing
# ════════════════════════════════════════════════════════════

def parse_txt(content: bytes) -> str:
    """Parse plain text / markdown with multi-encoding fallback."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def parse_pdf(content: bytes) -> str:
    """Parse PDF with per-page error isolation and dual-library fallback."""
    text_parts = []

    # Attempt 1: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as page_err:
                    logger.warning(f"pdfplumber: page {i} failed: {page_err}")
    except Exception as e:
        logger.warning(f"pdfplumber failed entirely, trying pypdf: {e}")
        text_parts = []

    # Attempt 2: pypdf (only if pdfplumber extracted nothing)
    if not text_parts:
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as page_err:
                    logger.warning(f"pypdf: page {i} failed: {page_err}")
        except Exception as e2:
            logger.error(f"Both PDF parsers failed: {e2}")
            return ""

    return "\n\n".join(text_parts)


def parse_docx(content: bytes) -> str:
    """Parse DOCX — extracts paragraphs AND table content."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))

        parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        # Tables (often contain structured data crucial for RAG)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts)
    except Exception as e:
        logger.error(f"DOCX parse failed: {e}")
        return ""


def parse_document(content: bytes, file_type: str) -> str:
    """Route to the correct parser based on file extension."""
    file_type = file_type.lower().strip(".")
    if file_type in ("txt", "md"):
        return parse_txt(content)
    elif file_type == "pdf":
        return parse_pdf(content)
    elif file_type == "docx":
        return parse_docx(content)
    else:
        return ""


# ════════════════════════════════════════════════════════════
# Text Chunking
# ════════════════════════════════════════════════════════════

def _split_sentences(text: str) -> list[str]:
    """
    Lightweight sentence splitter — Vietnamese friendly.

    Splits on ., !, ?, ; followed by whitespace, plus paragraph breaks.
    Does NOT need a model; good enough as the unit for semantic
    boundary detection (the embedding-cosine logic absorbs minor noise).
    """
    parts = re.split(r"(?<=[\.!?;])\s+|\n{2,}", text.strip())
    return [p.strip() for p in parts if p and p.strip()]


async def semantic_chunk_text(
    text: str,
    target_chars: int = CHUNK_CHARS,
    min_chars: int = MIN_CHUNK_LENGTH,
) -> list[str]:
    """
    Embedding-based semantic chunking.

    Embed each sentence, walk pairwise cosine distances, place a break
    wherever the distance crosses the configured percentile threshold
    OR the running chunk hits target_chars. Sentences within a chunk
    stay topically related; topic shifts produce a new chunk.

    On any error (embed failure, no sentences, etc.) we fall back to
    the deterministic chunker — this function NEVER raises, ingest
    must keep working.
    """
    if not text or not text.strip():
        return []
    if len(text) <= target_chars:
        return [text.strip()]

    try:
        sentences = _split_sentences(text)
        if len(sentences) <= 3:
            return chunk_text(text)

        # Embed all sentences with the dense embedder. Re-uses the same
        # batched embed_texts so retries / fallback already work.
        embs = await embed_texts(sentences, batch_size=64)
        if len(embs) != len(sentences):
            logger.warning(
                "semantic_chunk: embed count mismatch — falling back to rule-based"
            )
            return chunk_text(text)

        # Cosine distance between consecutive sentences. Embeddings come
        # back already L2-normalized (BGE-M3 returns unit vectors), so
        # dot product == cosine similarity.
        def _cos_dist(a: list[float], b: list[float]) -> float:
            s = 0.0
            for x, y in zip(a, b):
                s += x * y
            return 1.0 - s

        distances = [_cos_dist(embs[i], embs[i + 1]) for i in range(len(embs) - 1)]
        if not distances:
            return chunk_text(text)

        # Percentile threshold — distances above this mark a topic break.
        sorted_d = sorted(distances)
        pct = max(0.0, min(100.0, settings.RAG_SEMANTIC_THRESHOLD_PCT)) / 100.0
        idx = min(len(sorted_d) - 1, int(pct * (len(sorted_d) - 1)))
        threshold = sorted_d[idx]

        chunks: list[str] = []
        current: list[str] = []
        current_len = 0

        for i, sent in enumerate(sentences):
            current.append(sent)
            current_len += len(sent) + 1
            is_last = (i == len(sentences) - 1)
            # Decide whether to break AFTER this sentence
            break_now = False
            if is_last:
                break_now = True
            elif current_len >= target_chars:
                break_now = True
            elif distances[i] >= threshold and current_len >= min_chars:
                break_now = True

            if break_now:
                joined = " ".join(current).strip()
                if len(joined) >= min_chars:
                    chunks.append(joined)
                current = []
                current_len = 0

        if not chunks:
            return chunk_text(text)

        # Dedup as in rule-based path
        seen: set[str] = set()
        unique: list[str] = []
        for c in chunks:
            h = hashlib.md5(c.encode()).hexdigest()
            if h not in seen:
                seen.add(h)
                unique.append(c)
        return unique[:MAX_CHUNKS_PER_FILE]

    except Exception as e:
        logger.warning(f"semantic_chunk failed ({e}); falling back to rule-based")
        return chunk_text(text)


def chunk_text(
    text: str,
    chunk_size: int = CHUNK_CHARS,
    overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    """
    Split text into overlapping chunks on paragraph/sentence boundaries.
    Guards against: infinite loops, excessive chunk counts, duplicate chunks.
    """
    if not text or not text.strip():
        return []

    text = text.strip()

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    paragraphs = text.split("\n\n")

    current_chunk = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if current_chunk and len(current_chunk) + len(para) + 2 > chunk_size:
            chunks.append(current_chunk.strip())
            if overlap > 0 and len(current_chunk) > overlap:
                current_chunk = current_chunk[-overlap:] + "\n\n" + para
            else:
                current_chunk = para
        else:
            current_chunk = (current_chunk + "\n\n" + para) if current_chunk else para

        # Force-split oversized chunks — hard cap prevents infinite loops
        max_force_splits = 50
        split_count = 0
        while len(current_chunk) > chunk_size * 1.5 and split_count < max_force_splits:
            split_count += 1
            prev_len = len(current_chunk)

            split_at = chunk_size
            for sep in (". ", "! ", "? ", "\n", "; ", ", "):
                idx = current_chunk.rfind(sep, 0, chunk_size + 50)
                if idx > chunk_size * 0.5:
                    split_at = idx + len(sep)
                    break

            # Guarantee forward progress
            if split_at <= 0:
                split_at = chunk_size

            chunks.append(current_chunk[:split_at].strip())
            remainder = current_chunk[split_at:]

            if overlap > 0:
                overlap_start = max(0, split_at - overlap)
                tail = current_chunk[overlap_start:split_at]
                current_chunk = tail + remainder
            else:
                current_chunk = remainder

            # If we didn't actually shrink, force-break to avoid infinite loop
            if len(current_chunk) >= prev_len:
                logger.warning("Force-split not making progress, breaking out")
                break

        if len(chunks) >= MAX_CHUNKS_PER_FILE:
            logger.warning(f"Hit max chunk cap ({MAX_CHUNKS_PER_FILE}), truncating")
            break

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    # Filter tiny chunks + deduplicate
    seen = set()
    unique_chunks = []
    for c in chunks:
        if len(c) < MIN_CHUNK_LENGTH:
            continue
        h = hashlib.md5(c.encode()).hexdigest()
        if h not in seen:
            seen.add(h)
            unique_chunks.append(c)

    return unique_chunks[:MAX_CHUNKS_PER_FILE]


# ════════════════════════════════════════════════════════════
# Embedding Client — Local service with Managed API fallback
# ════════════════════════════════════════════════════════════
#
# Primary:   POST to settings.EMBED_HOST (local BGE-M3 microservice)
# Fallback:  POST to settings.EMBED_API_URL (OpenAI-compatible managed API)
#
# Auto-switch behavior: after LOCAL_FAILURE_THRESHOLD consecutive failures
# against the local service, we flip to the API for the rest of this process
# lifetime. The counter resets on the next successful local call.
# If no API URL is configured, we never switch — local errors propagate.
# ════════════════════════════════════════════════════════════

# Singleton HTTP clients — keep TCP connection pool alive across calls.
_embed_client: Optional[httpx.AsyncClient] = None
_rerank_client: Optional[httpx.AsyncClient] = None

# Failure tracking for auto-fallback
_local_embed_failures = 0
_local_rerank_failures = 0
_use_embed_api = False   # sticky: once switched, stay on API
_use_rerank_api = False


def get_embed_client() -> httpx.AsyncClient:
    """Return a reusable httpx client for the embed service."""
    global _embed_client
    if _embed_client is None or _embed_client.is_closed:
        _embed_client = httpx.AsyncClient(
            timeout=httpx.Timeout(180.0, connect=15.0),
            limits=httpx.Limits(
                max_keepalive_connections=10,
                max_connections=20,
                keepalive_expiry=30.0,
            ),
        )
    return _embed_client


def get_rerank_client() -> httpx.AsyncClient:
    """Return a reusable httpx client for the rerank service."""
    global _rerank_client
    if _rerank_client is None or _rerank_client.is_closed:
        _rerank_client = httpx.AsyncClient(
            timeout=httpx.Timeout(60.0, connect=10.0),
            limits=httpx.Limits(
                max_keepalive_connections=5,
                max_connections=10,
                keepalive_expiry=30.0,
            ),
        )
    return _rerank_client


async def close_http_clients():
    """Close singleton clients cleanly on app shutdown."""
    global _embed_client, _rerank_client
    if _embed_client is not None and not _embed_client.is_closed:
        await _embed_client.aclose()
        _embed_client = None
    if _rerank_client is not None and not _rerank_client.is_closed:
        await _rerank_client.aclose()
        _rerank_client = None


def _get_embed_hosts() -> list[str]:
    """
    Return the ordered list of local embed hosts to try.

    Primary = settings.EMBED_HOST. Fallbacks come from EMBED_HOST_FALLBACKS
    (comma-separated). Empties / dupes filtered. Order matters — chain
    stops at the first 200 response.
    """
    chain: list[str] = []
    seen: set[str] = set()

    primary = (settings.EMBED_HOST or "").strip().rstrip("/")
    if primary:
        chain.append(primary)
        seen.add(primary)

    fallbacks_raw = (settings.EMBED_HOST_FALLBACKS or "").strip()
    if fallbacks_raw:
        for h in fallbacks_raw.split(","):
            h = h.strip().rstrip("/")
            if h and h not in seen:
                chain.append(h)
                seen.add(h)
    return chain


async def _embed_batch_at_host(
    host: str, batch: list[str], max_retries: int = 2,
) -> list[list[float]]:
    """Call /embed on ONE host with retries. Raises on definitive failure."""
    client = get_embed_client()
    last_err: Exception | None = None
    for attempt in range(max_retries + 1):
        try:
            resp = await client.post(f"{host}/embed", json={"texts": batch})
            resp.raise_for_status()
            data = resp.json()
            batch_embs = data.get("embeddings", [])
            if len(batch_embs) != len(batch):
                raise ValueError(
                    f"count mismatch: sent {len(batch)}, got {len(batch_embs)}"
                )
            return batch_embs
        except (
            httpx.HTTPStatusError, httpx.ConnectError,
            httpx.ReadTimeout, httpx.RemoteProtocolError, ValueError,
        ) as e:
            last_err = e
            if attempt < max_retries:
                import asyncio
                await asyncio.sleep(2 ** attempt)
            else:
                raise RuntimeError(
                    f"Embed host {host} failed after {max_retries+1} attempts: {last_err}"
                ) from last_err
    # Unreachable, satisfies type checker
    raise RuntimeError(f"Embed host {host} failed: {last_err}")


async def _embed_via_local(
    texts: list[str], batch_size: int = 32, max_retries: int = 2,
) -> list[list[float]]:
    """
    Call local BGE-M3 microservices with multi-host fallback.

    Tries EMBED_HOST first, then each entry in EMBED_HOST_FALLBACKS in
    order. Per-batch granularity: if host A fails on batch 3, the
    remaining batches retry from host A first; we only escalate to
    host B if host A still fails. This keeps connection pools warm at
    the primary and avoids thrashing.

    Raises RuntimeError when ALL hosts fail on a single batch — the
    caller (`embed_texts`) catches that and triggers the API fallback.
    """
    hosts = _get_embed_hosts()
    if not hosts:
        raise RuntimeError("No EMBED_HOST configured")

    all_embeddings: list[list[float]] = []

    for batch_idx, i in enumerate(range(0, len(texts), batch_size)):
        batch = texts[i : i + batch_size]
        host_errors: list[str] = []
        batch_embs: list[list[float]] | None = None

        for host in hosts:
            try:
                batch_embs = await _embed_batch_at_host(host, batch, max_retries)
                if host != hosts[0]:
                    logger.info(
                        f"Embed batch {batch_idx} succeeded via fallback host {host}"
                    )
                break
            except Exception as e:
                host_errors.append(f"{host}: {e}")
                logger.warning(
                    f"Embed host {host} failed on batch {batch_idx} ({e}), "
                    f"trying next in chain"
                )
                continue

        if batch_embs is None:
            raise RuntimeError(
                f"All {len(hosts)} embed hosts failed on batch {batch_idx}. "
                + " | ".join(host_errors)
            )
        all_embeddings.extend(batch_embs)

    return all_embeddings


def _resolve_embed_api_key() -> str:
    """Use EMBED_API_KEY if set, else fall back to DASHSCOPE_API_KEY."""
    return (settings.EMBED_API_KEY or settings.DASHSCOPE_API_KEY or "").strip()


def _resolve_rerank_api_key() -> str:
    """Use RERANK_API_KEY if set, else fall back to DASHSCOPE_API_KEY."""
    return (settings.RERANK_API_KEY or settings.DASHSCOPE_API_KEY or "").strip()


def _build_resource_url(base_url: str, resource: str) -> str:
    """
    Append the resource path unless the URL already includes it.

    Lets users put either:
        EMBED_API_URL=https://.../v1            → we append /embeddings
        EMBED_API_URL=https://.../v1/embeddings → we use as-is
    """
    base = (base_url or "").rstrip("/")
    if not base:
        return ""
    leaf = base.rsplit("/", 1)[-1].lower()
    res_clean = resource.strip("/").lower()
    if leaf == res_clean:
        return base
    return f"{base}/{resource.strip('/')}"


async def _embed_via_api(
    texts: list[str], batch_size: int = 32,
) -> list[list[float]]:
    """
    Call OpenAI-compatible managed embedding API. Default target is
    Alibaba DashScope (text-embedding-v4); compatible with any provider
    exposing the OpenAI `/embeddings` schema.

    Payload sent: only standard OpenAI fields (model, input, dimensions,
    encoding_format). FPT-specific extensions (input_text_truncate,
    input_type) are dropped — DashScope rejects unknown params on some
    endpoints, and they were never load-bearing for retrieval quality.
    """
    all_embeddings = []
    client = get_embed_client()

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {_resolve_embed_api_key()}",
    }
    url = _build_resource_url(settings.EMBED_API_URL, "embeddings")

    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]

        payload = {
            "model": settings.EMBED_API_MODEL,
            "input": batch,
            "dimensions": settings.EMBED_API_DIMENSIONS,
            "encoding_format": "float",
        }

        resp = await client.post(url, json=payload, headers=headers)
        resp.raise_for_status()
        data = resp.json()

        # OpenAI-compatible response: data[{index, embedding}]
        items = sorted(data.get("data", []), key=lambda x: x.get("index", 0))
        embs = [it["embedding"] for it in items]
        if len(embs) != len(batch):
            raise ValueError(
                f"API returned {len(embs)} embeddings for {len(batch)} inputs"
            )
        all_embeddings.extend(embs)

    return all_embeddings


async def embed_texts(
    texts: list[str],
    batch_size: int = 32,
    max_retries: int = 2,
) -> list[list[float]]:
    """
    Embed texts. Routes based on settings.EMBED_MODE:
      "local"  → local service only; errors propagate
      "backup" → managed API only; errors propagate
      "auto"   → try local first; after LOCAL_FAILURE_THRESHOLD failures
                 in a row, sticky-switch to API for the rest of this process
    """
    global _local_embed_failures, _use_embed_api

    if not texts:
        return []

    mode = settings.EMBED_MODE.lower().strip()
    api_configured = bool(settings.EMBED_API_URL and _resolve_embed_api_key())
    local_configured = bool(_get_embed_hosts())

    # ── Mode: backup — go straight to API ──
    if mode == "backup":
        if not api_configured:
            raise RuntimeError(
                "EMBED_MODE=backup but EMBED_API_URL is empty or no API key "
                "(EMBED_API_KEY / DASHSCOPE_API_KEY) is set"
            )
        return await _embed_via_api(texts, batch_size=batch_size)

    # ── Mode: local — local only, no fallback ──
    if mode == "local":
        if not local_configured:
            # Misconfiguration guard: user set mode=local but no EMBED_HOST.
            # If API is configured, fall through to API instead of crashing.
            if api_configured:
                logger.warning(
                    "EMBED_MODE=local but EMBED_HOST empty — routing to API. "
                    "Set EMBED_MODE=backup explicitly to silence this warning."
                )
                return await _embed_via_api(texts, batch_size=batch_size)
            raise RuntimeError(
                "EMBED_MODE=local but EMBED_HOST is empty and no API configured"
            )
        return await _embed_via_local(texts, batch_size, max_retries)

    # ── Mode: auto — local with sticky fallback ──
    # Already in sticky API mode?
    if _use_embed_api and api_configured:
        return await _embed_via_api(texts, batch_size=batch_size)

    # Try local
    try:
        result = await _embed_via_local(texts, batch_size, max_retries)
        if _local_embed_failures > 0:
            logger.info("Local embed recovered, resetting failure counter")
            _local_embed_failures = 0
        return result
    except Exception as local_err:
        _local_embed_failures += 1
        logger.warning(
            f"Local embed failed ({_local_embed_failures}/"
            f"{settings.LOCAL_FAILURE_THRESHOLD}): {local_err}"
        )

        if not api_configured:
            raise  # no fallback to use

        over_threshold = _local_embed_failures >= settings.LOCAL_FAILURE_THRESHOLD
        if over_threshold:
            logger.warning(
                f"🔀 Sticky-switch to managed embed API: {settings.EMBED_API_URL}"
            )
            _use_embed_api = True

        # Retry this request via API
        try:
            return await _embed_via_api(texts, batch_size=batch_size)
        except Exception as api_err:
            raise RuntimeError(
                f"Both embed providers failed. Local: {local_err}. API: {api_err}"
            ) from api_err


async def embed_query(query: str) -> list[float]:
    """Embed a single query string."""
    result = await embed_texts([query])
    return result[0] if result else []


# ════════════════════════════════════════════════════════════
# Hybrid (dense + sparse) Embedding
# ════════════════════════════════════════════════════════════
#
# Hybrid retrieval requires BOTH dense and sparse vectors. Sparse is
# only produced by the local /embed_hybrid endpoint (backed by
# FlagEmbedding's BGEM3FlagModel). The managed API fallback only
# returns dense, so when hybrid is unavailable we transparently return
# empty sparse dicts and downstream code falls back to dense-only.
# ════════════════════════════════════════════════════════════

# Sticky flag: after we've discovered the local embed service doesn't
# support hybrid (older container without FlagEmbedding), stop probing.
_hybrid_probe_done = False
_hybrid_available = False


async def embed_texts_hybrid(
    texts: list[str], batch_size: int = 32,
) -> tuple[list[list[float]], list[dict[int, float]], bool]:
    """
    Return (dense, sparse, sparse_available).

    sparse_available=False means we couldn't get sparse weights (older
    embed service, API mode, etc.) — caller should fall back to dense
    only. sparse list will still be returned, but as empty dicts.
    """
    global _hybrid_probe_done, _hybrid_available

    if not texts:
        return [], [], _hybrid_available

    # Hybrid only works through the local service (sparse not in API)
    hosts = _get_embed_hosts()
    mode = settings.EMBED_MODE.lower().strip()

    if mode == "backup" or not hosts or _use_embed_api:
        # No local path → no sparse. Fall back to dense-only via existing path.
        dense = await embed_texts(texts, batch_size=batch_size)
        return dense, [{} for _ in texts], False

    client = get_embed_client()
    all_dense: list[list[float]] = []
    all_sparse: list[dict[int, float]] = []
    sparse_ok = True

    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        data: dict | None = None
        host_errors: list[str] = []

        # Try each host in chain. First 200 wins. Mirrors _embed_via_local
        # multi-host logic — see _get_embed_hosts() for ordering.
        for host in hosts:
            try:
                resp = await client.post(
                    f"{host}/embed_hybrid",
                    json={
                        "texts": batch,
                        "return_dense": True,
                        "return_sparse": True,
                    },
                )
                resp.raise_for_status()
                data = resp.json()
                if host != hosts[0]:
                    logger.info(
                        f"Hybrid embed batch {i // batch_size} via fallback host {host}"
                    )
                break
            except Exception as e:
                host_errors.append(f"{host}: {e}")
                logger.warning(
                    f"embed_hybrid host {host} failed ({e}), trying next"
                )
                continue

        if data is None:
            # All hosts down. Degrade to dense-only via the standard
            # embed_texts path (which itself walks the host chain and
            # falls to API as a last resort).
            logger.warning(
                f"All {len(hosts)} hybrid hosts failed; degrading to dense-only. "
                + " | ".join(host_errors)
            )
            _hybrid_probe_done = True
            _hybrid_available = False
            dense = await embed_texts(texts, batch_size=batch_size)
            return dense, [{} for _ in texts], False

        dense_batch = data.get("dense", [])
        sparse_batch_raw = data.get("sparse", [])
        if not data.get("sparse_available", False):
            sparse_ok = False

        all_dense.extend(dense_batch)
        # Cast string keys → int (Milvus sparse format)
        for sd in sparse_batch_raw:
            if not isinstance(sd, dict):
                all_sparse.append({})
                continue
            cast: dict[int, float] = {}
            for k, v in sd.items():
                try:
                    cast[int(k)] = float(v)
                except (TypeError, ValueError):
                    continue
            all_sparse.append(cast)

    # Pad sparse if length mismatch (defensive)
    while len(all_sparse) < len(all_dense):
        all_sparse.append({})

    _hybrid_probe_done = True
    _hybrid_available = sparse_ok
    return all_dense, all_sparse, sparse_ok


async def embed_query_hybrid(
    query: str,
) -> tuple[list[float], dict[int, float], bool]:
    """Hybrid embed a single query string."""
    dense, sparse, ok = await embed_texts_hybrid([query])
    if not dense:
        return [], {}, False
    return dense[0], (sparse[0] if sparse else {}), ok


# ════════════════════════════════════════════════════════════
# Milvus Client
# ════════════════════════════════════════════════════════════

EMBEDDING_DIM = 1024  # BGE-M3 output dimension

# Legacy collection name kept for backward compatibility. The active
# collection is read from settings.RAG_COLLECTION at call time, so
# flipping between v1/v2 is just an env change.
LEGACY_COLLECTION_NAME = "finhouse_chunks"


def _active_collection_name() -> str:
    """Resolve the active collection from settings (env-driven, no restart)."""
    name = (settings.RAG_COLLECTION or "").strip()
    return name or LEGACY_COLLECTION_NAME


def _is_v2_collection(name: str | None = None) -> bool:
    """v2 = collection that has sparse_embedding field + HNSW dense index."""
    n = name or _active_collection_name()
    return n != LEGACY_COLLECTION_NAME


# Track which collections we've already initialized this process. Lets us
# operate on both v1 and v2 in the same process (e.g. for migration).
_milvus_initialized: set[str] = set()
_milvus_lock = threading.Lock()


def _build_dense_index_params() -> dict:
    """Index params for the dense vector field."""
    idx_type = (settings.RAG_INDEX_TYPE or "HNSW").upper()
    if idx_type == "HNSW":
        return {
            "metric_type": "COSINE",
            "index_type": "HNSW",
            "params": {"M": 16, "efConstruction": 200},
        }
    # Legacy / fallback
    return {
        "metric_type": "COSINE",
        "index_type": "IVF_FLAT",
        "params": {"nlist": 128},
    }


def _dense_search_params() -> dict:
    """Search params matching the active index type."""
    idx_type = (settings.RAG_INDEX_TYPE or "HNSW").upper()
    if idx_type == "HNSW" and _is_v2_collection():
        return {"metric_type": "COSINE", "params": {"ef": 128}}
    return {"metric_type": "COSINE", "params": {"nprobe": 16}}


def _get_milvus_connection(collection_name: str | None = None):
    """
    Thread-safe Milvus connection + collection init.

    If collection_name is None, uses the active one from settings. The
    v2 schema adds a sparse_embedding field + SPARSE_INVERTED_INDEX so
    hybrid search works. v1 schema is left untouched.
    """
    from pymilvus import (
        connections, Collection, FieldSchema, CollectionSchema,
        DataType, utility,
    )

    name = collection_name or _active_collection_name()

    # Retry connect with exponential backoff. On a shared host with etcd
    # contention the gRPC handshake intermittently times out; one quick
    # hiccup shouldn't fail an entire ingest. We bump pymilvus's default
    # 10s connect timeout to 30s and retry up to 3 times (~total ~75s).
    import time as _time
    last_err: Exception | None = None
    for attempt in range(3):
        try:
            # Disconnect any stale alias before reconnecting (otherwise
            # pymilvus reuses the dead channel on next call)
            try:
                connections.disconnect("default")
            except Exception:
                pass
            connections.connect(
                alias="default",
                host=settings.MILVUS_HOST,
                port=str(settings.MILVUS_PORT),
                timeout=30,
            )
            last_err = None
            break
        except Exception as e:
            last_err = e
            wait = 2 ** attempt   # 1s, 2s, 4s
            logger.warning(
                f"Milvus connect attempt {attempt+1}/3 failed ({e}); "
                f"retrying in {wait}s"
            )
            _time.sleep(wait)
    if last_err is not None:
        logger.error(f"Milvus connection failed after 3 attempts: {last_err}")
        raise last_err

    with _milvus_lock:
        if name in _milvus_initialized:
            return

        if not utility.has_collection(name):
            v2 = _is_v2_collection(name)

            fields = [
                FieldSchema(
                    name="id", dtype=DataType.VARCHAR,
                    is_primary=True, max_length=128,
                ),
                FieldSchema(name="file_id", dtype=DataType.VARCHAR, max_length=64),
                FieldSchema(name="project_id", dtype=DataType.INT64),
                FieldSchema(name="file_name", dtype=DataType.VARCHAR, max_length=512),
                FieldSchema(name="chunk_index", dtype=DataType.INT64),
                FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=8192),
                FieldSchema(
                    name="embedding", dtype=DataType.FLOAT_VECTOR,
                    dim=EMBEDDING_DIM,
                ),
            ]
            if v2:
                fields.append(FieldSchema(
                    name="sparse_embedding",
                    dtype=DataType.SPARSE_FLOAT_VECTOR,
                ))

            schema = CollectionSchema(
                fields=fields,
                description=f"FinHouse document chunks ({'v2 hybrid' if v2 else 'v1 dense-only'})",
            )
            collection = Collection(name=name, schema=schema)

            collection.create_index(
                field_name="embedding",
                index_params=_build_dense_index_params(),
            )
            if v2:
                collection.create_index(
                    field_name="sparse_embedding",
                    index_params={
                        "index_type": "SPARSE_INVERTED_INDEX",
                        "metric_type": "IP",
                    },
                )
            logger.info(
                f"Created Milvus collection: {name} "
                f"(v2={v2}, dense_index={settings.RAG_INDEX_TYPE})"
            )
        _milvus_initialized.add(name)


def upsert_chunks(
    file_id: str,
    project_id: int,
    file_name: str,
    chunks: list[str],
    embeddings: list[list[float]],
    sparse_embeddings: Optional[list[dict[int, float]]] = None,
    collection_name: Optional[str] = None,
):
    """
    Insert chunks in batches to avoid Milvus payload size limits.

    sparse_embeddings: only used when writing into a v2 collection. List
    of {token_id: weight} dicts, one per chunk. If None on a v2 collection
    we insert empty dicts (chunk is dense-only retrievable).
    """
    from pymilvus import Collection

    if not chunks:
        return

    name = collection_name or _active_collection_name()
    _get_milvus_connection(name)
    collection = Collection(name)
    v2 = _is_v2_collection(name)

    INSERT_BATCH = 500
    total_inserted = 0

    for start in range(0, len(chunks), INSERT_BATCH):
        end = min(start + INSERT_BATCH, len(chunks))
        batch_chunks = chunks[start:end]
        batch_embs = embeddings[start:end]

        ids = [f"{file_id}_{i}" for i in range(start, end)]
        file_ids = [str(file_id)] * len(batch_chunks)
        project_ids = [int(project_id)] * len(batch_chunks)
        file_names = [file_name[:512]] * len(batch_chunks)
        chunk_indices = list(range(start, end))
        texts = [c[:8000] for c in batch_chunks]

        row_data = [
            ids, file_ids, project_ids, file_names,
            chunk_indices, texts, batch_embs,
        ]
        if v2:
            if sparse_embeddings is not None:
                batch_sparse = sparse_embeddings[start:end]
            else:
                batch_sparse = [None] * len(batch_chunks)
            # Milvus SPARSE_FLOAT_VECTOR rejects empty / None — every row
            # must be a non-empty {token_id: weight} dict. When sparse
            # isn't available (local embed down → API fallback, which
            # only returns dense), we insert a degenerate placeholder
            # {0: 1e-9}. Sparse search on such a vector contributes ~0
            # to RRF, so the chunk stays retrievable via dense and
            # ranking isn't skewed. Empty dicts from FlagEmbedding for
            # chunks with no recognized tokens get the same treatment.
            batch_sparse = [
                (sv if isinstance(sv, dict) and sv else {0: 1e-9})
                for sv in batch_sparse
            ]
            row_data.append(batch_sparse)

        collection.insert(row_data)
        total_inserted += len(batch_chunks)

    collection.flush()
    logger.info(
        f"Inserted {total_inserted} chunks for {file_name} "
        f"(file_id={file_id}, collection={name})"
    )


def delete_file_chunks(file_id: str, collection_name: Optional[str] = None):
    """Remove all chunks for a file from Milvus. Non-fatal on error."""
    from pymilvus import Collection

    name = collection_name or _active_collection_name()
    try:
        _get_milvus_connection(name)
        collection = Collection(name)
        collection.load()
        collection.delete(f'file_id == "{file_id}"')
        collection.flush()
        logger.info(f"Deleted chunks for file_id={file_id} (collection={name})")
    except Exception as e:
        logger.warning(f"delete_file_chunks({file_id}) failed (non-fatal): {e}")


# Whitelist for ticker-style filename prefixes. VN tickers are
# alphanumeric uppercase, ≤ 8 chars. Anything outside is dropped.
_PREFIX_OK_RE = re.compile(r"^[A-Za-z0-9]+_?$")

# Multiplier applied to a chunk's vector score when its file name
# starts with a known ticker prefix (e.g. "ACB_"). Boost is
# intentionally modest: it nudges the ranking toward
# convention-named files but never drowns out a strongly relevant
# user-uploaded file that doesn't follow the convention.
FILENAME_BOOST_MULT = 1.20


def _normalize_prefixes(prefixes: list[str]) -> list[str]:
    """Uppercase, dedupe, and ensure each prefix ends with `_`."""
    out = []
    seen = set()
    for p in prefixes or []:
        p = (p or "").strip().upper()
        if not p or len(p) > 32 or not _PREFIX_OK_RE.match(p):
            continue
        if not p.endswith("_"):
            p = p + "_"
        if p in seen:
            continue
        seen.add(p)
        out.append(p)
    return out


def _project_scope_expr(project_id: int) -> str:
    """Build the Milvus filter expression for the project-scope rules."""
    if project_id < 0:
        return f"project_id == {project_id}"
    if project_id == 0:
        return "project_id == 0"
    return f"project_id in [0, {project_id}]"


def _materialize_hits(
    result_iter, safe_prefixes: list[str], score_field: str = "score",
) -> list[dict]:
    """Convert Milvus search result into hit dicts with optional boost."""
    hits: list[dict] = []
    for hit_list in result_iter:
        for hit in hit_list:
            fname = hit.entity.get("file_name") or ""
            raw_score = float(hit.score)
            score = raw_score
            boosted = False
            if safe_prefixes:
                fname_upper = fname.upper()
                if any(fname_upper.startswith(p) for p in safe_prefixes):
                    score = raw_score * FILENAME_BOOST_MULT
                    boosted = True
            hits.append({
                "id": hit.id,
                score_field: score,
                f"raw_{score_field}": raw_score,
                "filename_boosted": boosted,
                "file_id": hit.entity.get("file_id"),
                "file_name": fname,
                "chunk_index": hit.entity.get("chunk_index"),
                "text": hit.entity.get("text"),
                "project_id": hit.entity.get("project_id"),
            })
    return hits


async def search_chunks(
    query_embedding: list[float],
    project_id: int,
    top_k: int = 20,
    file_name_prefixes: Optional[list[str]] = None,
) -> list[dict]:
    """
    Dense-only search (legacy + fallback). For hybrid search use
    `hybrid_search_chunks`.

    Scope rules:
      • project_id == 0 (Inbox) — BASE KNOWLEDGE, accessible to all users.
      • Any other project_id: that project + base (0).
      • Negative project_id (incognito): only that project, not base.
    """
    from pymilvus import Collection

    name = _active_collection_name()
    _get_milvus_connection(name)
    collection = Collection(name)

    try:
        collection.load()
    except Exception as e:
        logger.debug(f"collection.load() note (may already be loaded): {e}")

    expr = _project_scope_expr(project_id)
    safe_prefixes = _normalize_prefixes(file_name_prefixes or [])
    effective_limit = top_k * 2 if safe_prefixes else top_k

    results = collection.search(
        data=[query_embedding],
        anns_field="embedding",
        param=_dense_search_params(),
        limit=effective_limit,
        expr=expr,
        output_fields=["file_id", "file_name", "chunk_index", "text", "project_id"],
    )

    hits = _materialize_hits(results, safe_prefixes, score_field="score")

    if safe_prefixes:
        hits.sort(key=lambda h: h["score"], reverse=True)
        hits = hits[:top_k]

    boosted_count = sum(1 for h in hits if h.get("filename_boosted"))
    logger.info(
        f"Milvus dense search: collection={name} project={project_id} "
        f"prefixes={safe_prefixes or '-'} → {len(hits)} hits "
        f"({boosted_count} boosted)"
    )
    return hits


async def hybrid_search_chunks(
    query_dense_list: list[list[float]],
    query_sparse_list: list[dict[int, float]],
    project_id: int,
    top_k: int = 20,
    file_name_prefixes: Optional[list[str]] = None,
) -> list[dict]:
    """
    Multi-query hybrid search with RRF fusion.

    Runs N dense searches AND (if available) N sparse searches against
    the active v2 collection, then fuses all ranked lists via Reciprocal
    Rank Fusion. Filename-prefix boost is applied per-list before RRF.

    Falls back to dense-only fusion if:
      • Collection is v1 (no sparse field)
      • RAG_HYBRID_ENABLED is False
      • Any sparse vector is empty
    """
    from pymilvus import Collection

    if not query_dense_list:
        return []

    name = _active_collection_name()
    _get_milvus_connection(name)
    collection = Collection(name)
    try:
        collection.load()
    except Exception as e:
        logger.debug(f"collection.load() note: {e}")

    v2 = _is_v2_collection(name)
    use_sparse = (
        v2
        and settings.RAG_HYBRID_ENABLED
        and query_sparse_list
        and any(sv for sv in query_sparse_list)
    )

    expr = _project_scope_expr(project_id)
    safe_prefixes = _normalize_prefixes(file_name_prefixes or [])
    # Pull a wider pool per list — RRF needs ranked depth to work well.
    per_list_limit = max(top_k * 2, 40)
    output_fields = ["file_id", "file_name", "chunk_index", "text", "project_id"]

    ranked_lists: list[list[dict]] = []

    # Dense queries: one search per query (rewrite + HyDE passages)
    for q_dense in query_dense_list:
        if not q_dense:
            continue
        res = collection.search(
            data=[q_dense],
            anns_field="embedding",
            param=_dense_search_params(),
            limit=per_list_limit,
            expr=expr,
            output_fields=output_fields,
        )
        hits = _materialize_hits(res, safe_prefixes, score_field="score")
        if safe_prefixes:
            hits.sort(key=lambda h: h["score"], reverse=True)
        if hits:
            ranked_lists.append(hits)

    # Sparse queries: only when v2 + hybrid enabled + non-empty vectors
    if use_sparse:
        for q_sparse in query_sparse_list:
            if not q_sparse:
                continue
            try:
                res = collection.search(
                    data=[q_sparse],
                    anns_field="sparse_embedding",
                    param={"metric_type": "IP", "params": {}},
                    limit=per_list_limit,
                    expr=expr,
                    output_fields=output_fields,
                )
            except Exception as e:
                logger.warning(f"sparse search failed (skipping): {e}")
                break
            hits = _materialize_hits(res, safe_prefixes, score_field="score")
            if safe_prefixes:
                hits.sort(key=lambda h: h["score"], reverse=True)
            if hits:
                ranked_lists.append(hits)

    if not ranked_lists:
        return []

    fused = _rrf_fuse(ranked_lists, k=settings.RAG_RRF_K)
    fused = fused[:top_k]

    logger.info(
        f"Hybrid search: collection={name} project={project_id} "
        f"n_lists={len(ranked_lists)} (sparse_used={use_sparse}) "
        f"prefixes={safe_prefixes or '-'} → {len(fused)} fused"
    )
    return fused


def _rrf_fuse(
    ranked_lists: list[list[dict]], k: int = 60,
) -> list[dict]:
    """
    Reciprocal Rank Fusion.

    score(d) = Σ 1 / (k + rank_in_list_i(d))

    Deduplicates by chunk `id`. Filename-boosted flag survives if any
    contributing list had it boosted. Keeps the highest raw score across
    lists for diagnostics. RRF score is written under `score` so downstream
    threshold + rerank still works unchanged.
    """
    scores: dict[str, float] = {}
    keep: dict[str, dict] = {}
    boosted_in: dict[str, bool] = {}
    best_raw: dict[str, float] = {}

    for ranked in ranked_lists:
        for rank, hit in enumerate(ranked):
            hid = hit.get("id") or f"{hit.get('file_id')}_{hit.get('chunk_index')}"
            scores[hid] = scores.get(hid, 0.0) + 1.0 / (k + rank + 1)
            if hid not in keep:
                keep[hid] = hit
            boosted_in[hid] = boosted_in.get(hid, False) or hit.get("filename_boosted", False)
            best_raw[hid] = max(best_raw.get(hid, 0.0), float(hit.get("score", 0)))

    out: list[dict] = []
    for hid, rrf in sorted(scores.items(), key=lambda x: x[1], reverse=True):
        h = dict(keep[hid])
        h["score"] = rrf                    # primary score = RRF
        h["raw_score"] = best_raw.get(hid, 0.0)
        h["filename_boosted"] = boosted_in.get(hid, False)
        out.append(h)
    return out


# ════════════════════════════════════════════════════════════
# Reranker Client
# ════════════════════════════════════════════════════════════

# ════════════════════════════════════════════════════════════
# Reranker Client — Local service with Managed API fallback
# ════════════════════════════════════════════════════════════

def _get_rerank_hosts() -> list[str]:
    """
    Return the ordered list of local rerank hosts to try.

    Primary = settings.RERANK_HOST. Fallbacks come from RERANK_HOST_FALLBACKS
    (comma-separated). Empties / dupes filtered. Order matters — chain
    stops at the first successful response.
    """
    chain: list[str] = []
    seen: set[str] = set()

    primary = (settings.RERANK_HOST or "").strip().rstrip("/")
    if primary:
        chain.append(primary)
        seen.add(primary)

    fallbacks_raw = (settings.RERANK_HOST_FALLBACKS or "").strip()
    if fallbacks_raw:
        for h in fallbacks_raw.split(","):
            h = h.strip().rstrip("/")
            if h and h not in seen:
                chain.append(h)
                seen.add(h)
    return chain


async def _rerank_at_host(
    host: str, query: str, documents: list[str], top_n: int,
) -> list[dict]:
    """Call /rerank on ONE host. Raises on failure."""
    client = get_rerank_client()
    resp = await client.post(
        f"{host}/rerank",
        json={"query": query, "documents": documents, "top_n": top_n},
    )
    resp.raise_for_status()
    data = resp.json()
    return [
        {"index": r["index"], "score": float(r["score"])}
        for r in data.get("results", [])
    ]


async def _rerank_via_local(
    query: str, documents: list[str], top_n: int,
) -> list[dict]:
    """
    Call local rerank microservices with multi-host fallback.

    Tries RERANK_HOST first, then each entry in RERANK_HOST_FALLBACKS
    in order. First 200 wins. Raises RuntimeError when ALL hosts fail —
    caller (`rerank_chunks`) then triggers the API fallback.
    """
    hosts = _get_rerank_hosts()
    if not hosts:
        raise RuntimeError("No RERANK_HOST configured")

    host_errors: list[str] = []
    for host in hosts:
        try:
            results = await _rerank_at_host(host, query, documents, top_n)
            if host != hosts[0]:
                logger.info(f"Rerank succeeded via fallback host {host}")
            return results
        except Exception as e:
            host_errors.append(f"{host}: {e}")
            logger.warning(
                f"Rerank host {host} failed ({e}), trying next in chain"
            )
            continue

    raise RuntimeError(
        f"All {len(hosts)} rerank hosts failed. " + " | ".join(host_errors)
    )


async def _rerank_via_api(
    query: str, documents: list[str], top_n: int,
) -> list[dict]:
    """
    Call managed rerank API. Default target is Alibaba DashScope
    (qwen3-rerank). Response shape is the standard rerank schema:
    {results: [{index, relevance_score}]}.

    Note the resource path is `/reranks` (plural) per DashScope.
    `_build_resource_url` tolerates either base URL form.
    """
    client = get_rerank_client()
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {_resolve_rerank_api_key()}",
    }
    url = _build_resource_url(settings.RERANK_API_URL, "reranks")

    payload = {
        "model": settings.RERANK_API_MODEL,
        "query": query,
        "documents": documents,
        "top_n": top_n,
    }

    resp = await client.post(url, json=payload, headers=headers)
    resp.raise_for_status()
    data = resp.json()
    return [
        {"index": r["index"], "score": float(r.get("relevance_score", r.get("score", 0)))}
        for r in data.get("results", [])
    ]


async def rerank_chunks(
    query: str, chunks: list[dict], top_n: int = 5,
) -> list[dict]:
    """
    Rerank with auto-fallback from local service to managed API.
    Graceful degradation: if both fail, returns chunks in original order.
    """
    global _local_rerank_failures, _use_rerank_api

    if not chunks:
        return []

    valid_chunks = [c for c in chunks if c.get("text", "").strip()]
    if not valid_chunks:
        return []

    documents = [c["text"][:2000] for c in valid_chunks]

    async def _attach_scores(rerank_results: list[dict]) -> list[dict]:
        reranked = []
        for r in rerank_results:
            idx = r["index"]
            if 0 <= idx < len(valid_chunks):
                chunk = valid_chunks[idx].copy()
                chunk["rerank_score"] = r["score"]
                reranked.append(chunk)
        return reranked

    mode = settings.RERANK_MODE.lower().strip()
    api_configured = bool(settings.RERANK_API_URL and _resolve_rerank_api_key())
    local_configured = bool(_get_rerank_hosts())

    # Rerank is non-critical — graceful fallback to raw order on error
    def _raw_order_fallback(why: str) -> list[dict]:
        logger.info(f"Returning raw order: {why}")
        return valid_chunks[:top_n]

    # ── Mode: backup — API only ──
    if mode == "backup":
        if not api_configured:
            return _raw_order_fallback("RERANK_MODE=backup but API not configured")
        try:
            return await _attach_scores(
                await _rerank_via_api(query, documents, top_n)
            )
        except Exception as e:
            return _raw_order_fallback(f"API rerank failed: {e}")

    # ── Mode: local — local only ──
    if mode == "local":
        if not local_configured:
            # Misconfiguration guard: silently route to API if available
            if api_configured:
                logger.warning(
                    "RERANK_MODE=local but RERANK_HOST empty — routing to API"
                )
                try:
                    return await _attach_scores(
                        await _rerank_via_api(query, documents, top_n)
                    )
                except Exception as e:
                    return _raw_order_fallback(f"API rerank failed: {e}")
            return _raw_order_fallback("No rerank provider configured")
        try:
            return await _attach_scores(
                await _rerank_via_local(query, documents, top_n)
            )
        except Exception as e:
            return _raw_order_fallback(f"Local rerank failed: {e}")

    # ── Mode: auto ──
    if _use_rerank_api and api_configured:
        try:
            return await _attach_scores(
                await _rerank_via_api(query, documents, top_n)
            )
        except Exception as e:
            return _raw_order_fallback(f"Sticky API rerank failed: {e}")

    # Try local
    try:
        results = await _rerank_via_local(query, documents, top_n)
        if _local_rerank_failures > 0:
            logger.info("Local rerank recovered, resetting failure counter")
            _local_rerank_failures = 0
        return await _attach_scores(results)

    except Exception as local_err:
        _local_rerank_failures += 1
        logger.warning(
            f"Local rerank failed ({_local_rerank_failures}/"
            f"{settings.LOCAL_FAILURE_THRESHOLD}): {local_err}"
        )

        if not api_configured:
            return _raw_order_fallback("No API configured")

        over_threshold = _local_rerank_failures >= settings.LOCAL_FAILURE_THRESHOLD
        if over_threshold:
            logger.warning("🔀 Sticky-switch to managed rerank API")
            _use_rerank_api = True

        try:
            return await _attach_scores(
                await _rerank_via_api(query, documents, top_n)
            )
        except Exception as api_err:
            return _raw_order_fallback(
                f"Both rerank providers failed. Local: {local_err}. API: {api_err}"
            )


# ════════════════════════════════════════════════════════════
# Full Ingest Pipeline
# ════════════════════════════════════════════════════════════

async def ingest_file(
    file_id: str,
    file_content: bytes,
    file_name: str,
    file_type: str,
    project_id: int,
    update_status_callback=None,
) -> dict:
    """
    Full ingest pipeline: parse → chunk → embed → Milvus upsert.
    Returns {status, chunks_count, error?}.
    """
    result = {"status": "ready", "chunks_count": 0, "error": None}

    # 1. Validate format
    ft = file_type.lower().strip(".")
    if ft not in SUPPORTED_EXTENSIONS:
        result["status"] = "failed"
        result["error"] = f"Unsupported file type: {file_type}"
        return result

    # 2. Size guard
    size_mb = len(file_content) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        result["status"] = "failed"
        result["error"] = (
            f"File too large: {size_mb:.1f} MB (max {MAX_FILE_SIZE_MB} MB)"
        )
        return result

    try:
        # 3. Parse
        logger.info(f"Parsing {file_name} ({ft}, {size_mb:.1f} MB)...")
        text = parse_document(file_content, ft)
        if not text or len(text.strip()) < 10:
            result["status"] = "failed"
            result["error"] = "No extractable text found in document"
            return result

        logger.info(f"Extracted {len(text)} chars from {file_name}")

        # 4. Chunk — semantic if enabled (slower), otherwise rule-based
        if settings.RAG_SEMANTIC_CHUNKING:
            logger.info(f"Semantic chunking {file_name}...")
            chunks = await semantic_chunk_text(text)
        else:
            chunks = chunk_text(text)
        if not chunks:
            result["status"] = "failed"
            result["error"] = "Document produced no usable text chunks"
            return result

        logger.info(f"Split into {len(chunks)} chunks")

        # 5. Embed — hybrid (dense + sparse) when v2 + hybrid enabled,
        # otherwise dense only. embed_texts_hybrid handles both paths
        # and reports sparse availability.
        active_collection = _active_collection_name()
        write_sparse = (
            _is_v2_collection(active_collection)
            and settings.RAG_HYBRID_ENABLED
        )

        logger.info(
            f"Embedding {len(chunks)} chunks "
            f"(hybrid={write_sparse}, collection={active_collection})..."
        )
        if write_sparse:
            embeddings, sparse_embs, sparse_ok = await embed_texts_hybrid(chunks)
            if not sparse_ok:
                # Local hybrid not available — degrade to dense-only insert
                sparse_embs = None
        else:
            embeddings = await embed_texts(chunks)
            sparse_embs = None

        if len(embeddings) != len(chunks):
            result["status"] = "failed"
            result["error"] = (
                f"Embedding count mismatch: {len(chunks)} chunks "
                f"but {len(embeddings)} embeddings"
            )
            return result

        # 6. Delete old chunks in the ACTIVE collection only (idempotent
        # re-ingest). We deliberately do NOT touch the legacy collection
        # — keeping v1 untouched is the rollback guarantee. If you want
        # to also clean up v1 after a successful v2 migration, run a
        # separate dedicated cleanup pass.
        delete_file_chunks(file_id, collection_name=active_collection)

        # 7. Upsert into active Milvus collection
        logger.info(f"Upserting into Milvus collection={active_collection}...")
        upsert_chunks(
            file_id=file_id,
            project_id=project_id,
            file_name=file_name,
            chunks=chunks,
            embeddings=embeddings,
            sparse_embeddings=sparse_embs,
            collection_name=active_collection,
        )

        result["chunks_count"] = len(chunks)
        logger.info(f"✅ Ingest complete: {file_name} → {len(chunks)} chunks")

    except Exception as e:
        logger.error(f"Ingest failed for {file_name}: {e}", exc_info=True)
        result["status"] = "failed"
        result["error"] = str(e)[:500]

    return result


# ════════════════════════════════════════════════════════════
# RAG Retrieval (used by chat router)
# ════════════════════════════════════════════════════════════

MIN_RELEVANCE_SCORE = 0.3


async def retrieve_context(
    query: str | list[str],
    project_id: int,
    top_k: int = 20,
    top_n_rerank: int = 5,
    file_name_prefixes: Optional[list[str]] = None,
    rerank_query: Optional[str] = None,
) -> list[dict]:
    """
    Full RAG retrieval: embed → (hybrid) search → fuse → rerank → top.

    `query` accepts:
      • str — single query (legacy behavior)
      • list[str] — multiple queries (HyDE: rewritten + hypothetical
        passages). All are searched in parallel and RRF-fused.

    `rerank_query` overrides which query string the cross-encoder uses
    for reranking. With HyDE we want the cross-encoder scoring against
    the user's *real* question, not the synthetic hypothetical passages.
    If None, defaults to the first item of `query`.

    Path selection:
      • v2 collection + RAG_HYBRID_ENABLED → hybrid_search_chunks (RRF
        across dense + sparse + multi-query)
      • Otherwise → legacy dense-only search_chunks (RRF only when
        multiple queries given)

    Filename-prefix boost: see search_chunks for the rationale.
    """
    queries: list[str] = [query] if isinstance(query, str) else [
        q for q in (query or []) if q and q.strip()
    ]
    if not queries:
        return []

    # Cap fan-out — don't run unlimited searches
    queries = queries[: max(1, settings.RAG_HYDE_N_PASSAGES + 1)]
    rerank_q = rerank_query or queries[0]

    try:
        active_collection = _active_collection_name()
        use_hybrid_path = (
            _is_v2_collection(active_collection)
            and settings.RAG_HYBRID_ENABLED
        )

        if use_hybrid_path:
            dense_vecs, sparse_vecs, _ok = await embed_texts_hybrid(queries)
            candidates = await hybrid_search_chunks(
                query_dense_list=dense_vecs,
                query_sparse_list=sparse_vecs,
                project_id=project_id,
                top_k=top_k,
                file_name_prefixes=file_name_prefixes,
            )
        else:
            # Legacy / v1 dense-only. With multi-query, RRF across the
            # per-query dense lists (no sparse component).
            dense_vecs = await embed_texts(queries)
            if not dense_vecs:
                return []
            if len(dense_vecs) == 1:
                candidates = await search_chunks(
                    dense_vecs[0], project_id, top_k=top_k,
                    file_name_prefixes=file_name_prefixes,
                )
            else:
                # Multi-query on v1: run dense per query, RRF in-app
                ranked: list[list[dict]] = []
                for vec in dense_vecs:
                    if not vec:
                        continue
                    r = await search_chunks(
                        vec, project_id, top_k=top_k * 2,
                        file_name_prefixes=file_name_prefixes,
                    )
                    if r:
                        ranked.append(r)
                candidates = _rrf_fuse(ranked, k=settings.RAG_RRF_K)[:top_k] if ranked else []

        if not candidates:
            return []

        # Score-threshold filter. NOTE: with RRF active, "score" is a
        # 1/(k+rank) value (~0.01–0.03 range), not cosine similarity.
        # We compute an RRF-aware floor instead of the cosine 0.3 cutoff.
        if use_hybrid_path or len(queries) > 1:
            # RRF score threshold: pass through everything that hit at
            # least one ranked position. Reranker handles the real cut.
            relevant = [c for c in candidates if c.get("score", 0) > 0]
        else:
            relevant = [
                c for c in candidates
                if c.get("score", 0) >= MIN_RELEVANCE_SCORE
            ]

        if not relevant:
            logger.info(
                f"All {len(candidates)} candidates below relevance threshold"
            )
            return []

        # Rerank with the user-facing question, NOT the HyDE passages —
        # the cross-encoder must score against the real intent.
        reranked = await rerank_chunks(rerank_q, relevant, top_n=top_n_rerank)
        return reranked

    except Exception as e:
        logger.warning(f"RAG retrieval failed: {e}")
        return []
